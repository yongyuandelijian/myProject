from docx import Document

doc = Document(r"..\外部信息交换系统数据字典.docx")
# paragraphs=doc.paragraphs
tables = doc.tables

# 标题
# for paragraph in paragraphs:
#     print(type(paragraph))
#     print(paragraph.text)

# 获取表格
tc = 0
temp_bxx = []  # 用来存放表信息
for table in tables:
    tc += 1
    print("================== 第{tc}个表 =======================".format(tc=tc))
    if tc > 2 and tc % 2 == 1:
        temp_bxx = []
    if tc % 2 == 1:
        cc = 0
        for col in table.columns:
            cc += 1
            # print("================== 第{cc}列 =======================".format(cc=cc))
            if cc == 2:
                for cell in col.cells:
                    temp_bxx.append(cell.text)
    elif tc % 2 == 0:
        rc = 0
        for row in table.rows:
            rc += 1
            # print("================== 第{rc}行 =======================".format(rc=rc))
            temp_lxx = []  # 用来存放列基本信息
            for xx in temp_bxx:
                temp_lxx.append(xx)
            for cell in row.cells:
                if rc != 1:
                    temp_lxx.append(cell.text)
            print("要插入的列数据是：", temp_lxx)