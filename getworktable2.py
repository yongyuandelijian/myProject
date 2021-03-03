# 功能： 从word中获取到表注释信息，然后在生成数据字典
# aaa 20201211
# 难点在于如何将文字中的表明提取出来，和他自己的列注释去对应起来，然后在将结果插入Excel
from docx import Document
import re
import xlsxwriter

doc = Document(r"金税三期社保费数据字典.docx")
# paragraphs=doc.paragraphs


def hqsj():
    tables = doc.tables
    print(len(tables))
    # 获取表格
    tc = 0
    temp_bxx = []  # 用来存放表信息
    for table in tables:
        tc += 1
        print("================== 第{tc}个表 =======================".format(tc=tc))
        if tc > 2 and tc % 2 == 1:
            temp_bxx = []
        if tc % 2 == 1:
            cc = 0
            for col in table.columns:
                cc += 1
                # print("================== 第{cc}列 =======================".format(cc=cc))
                if cc == 2:
                    for cell in col.cells:
                        temp_bxx.append(cell.text)
        elif tc % 2 == 0:
            rc = 0
            for row in table.rows:
                rc += 1
                # print("================== 第{rc}行 =======================".format(rc=rc))
                temp_lxx = []  # 用来存放列基本信息
                for xx in temp_bxx:
                    temp_lxx.append(xx)
                for cell in row.cells:
                    if rc != 1:
                        temp_lxx.append(cell.text)
                print("要插入的列数据是：", temp_lxx)


def getdata():
    tables=doc.tables
    # 清理出我们需要的那部分标题,并且去除重复标题，每个表格对应一个标题，然后将数据换算成我们需要的表名称
    # qlwz=set()   实际上并不行，因为集合是无序的，存储进去之后顺序全部乱了
    qlwz=[]
    for dl in doc.paragraphs:
        wz=dl.text
        sl=len(re.findall('的列清单',wz))
        if sl>0:
            wz=re.findall('[a-zA-Z0-9_]+',wz)
            if wz:
                qlwz.append(wz[0])
    # print(qlwz)
    print("word中存在的表格是{bgs}个，段落数是{dls}个".format(bgs=len(tables), dls=len(qlwz)))

    # ========================= 开始存储数据 ========================= #
    workbook = xlsxwriter.Workbook(r'金税三期社保费数据字典.xlsx')
    bt = workbook.add_format(
        {'bold': True, 'bg_color': '9BC2E6', 'align': 'centre', 'valign': 'vcentre',
         'border': 1})
    nr = workbook.add_format(
        {'align': 'centre', 'valign': 'vcentre',
         'border': 1})
    # 然后来做一对一拼接，做好数据
    worksheet1 = workbook.add_worksheet("表清单注释")
    worksheet1.write(0, 0, '表名称',bt)
    worksheet1.write(0, 1, '表注释', bt)
    worksheet1.set_column(0, 0, 50)
    worksheet1.set_column(1, 1, 50)
    worksheet2=workbook.add_worksheet("数据表列注释")
    worksheet2.write(0, 0, '表名称', bt)
    worksheet2.write(0, 1, '列名称', bt)
    worksheet2.write(0, 2, '列注释', bt)
    worksheet2.write(0, 3, '数据类型', bt)
    worksheet2.write(0, 4, '是否主键', bt)
    worksheet2.write(0, 5, '是否可为空', bt)
    worksheet2.set_column(0, 3, 25)
    worksheet2.set_column(4, 5, 15)
    rownum = 1
    for i in range(len(tables)):
        if i==2:
            rownum = 1
        # 如果是第一个表则是表清单，将内容打印在第一个sheet页面，如果是其他表，则需要增加表名那一列,存储在第二个sheet页面
        if i==0:
            rownum=1
            for row in tables[i].rows:
                colnum = 0
                for cell in row.cells:
                    # print('现在是第{rownum}行，{colnum}列'.format(rownum=rownum,colnum=colnum))
                    # print(cell.text)
                    # 存储在第一个sheet页
                    worksheet1.write(rownum,colnum,cell.text)
                    colnum+=1
                rownum += 1
        else:
            # print(i)
            table_row=0
            for row in tables[i].rows:
                if table_row==0:
                    table_row += 1
                    continue
                colnum=1
                for cell in row.cells:
                    # print(cell.text)
                    # print(qlwz[i-1])
                    # 存在第二个sheet页
                    worksheet2.write(rownum, 0, qlwz[i-1])
                    worksheet2.write(rownum,colnum,cell.text)
                    colnum+=1
                rownum += 1
                table_row+=1

    workbook.close()




if __name__ == '__main__':
    getdata()