from docx import Document
import re

def getworktable():
    doc = Document(r"..\外部信息交换系统数据字典.docx")
    # paragraphs=doc.paragraphs
    tables = doc.tables

    # 标题
    # for paragraph in paragraphs:
    #     print(type(paragraph))
    #     print(paragraph.text)
    # 声明一个函数，将sql直接写入文件
    def xrwj(text):
        with open("D:\sql.txt") as f:
            pass

    # 获取表格
    tc = 0
    temp_bxx = []  # 用来存放表信息
    for table in tables:
        tc += 1
        # print("================== 第{tc}个表 =======================".format(tc=tc))
        if tc > 2 and tc % 2 == 1:
            temp_bxx = []
        if tc % 2 == 1:
            cc = 0
            for col in table.columns:
                cc += 1
                # print("================== 第{cc}列 =======================".format(cc=cc))
                if cc == 2:
                    for cell in col.cells:
                        wb = cell.text.replace('\n', '')
                        temp_bxx.append(wb)
                    temp_altertablecommentsql = "alter table sc_jc_gszj.cx_ypt_sjcc_{table_name} set comment '{table_comment}';".format(
                        table_name=temp_bxx[1], table_comment=temp_bxx[2])
                    temp_altertablecommentsql = temp_altertablecommentsql.upper()
                    if re.findall('WBJH', temp_altertablecommentsql):
                        print(temp_altertablecommentsql)
        elif tc % 2 == 0:
            rc = 0
            for row in table.rows:
                rc += 1
                # print("================== 第{rc}行 =======================".format(rc=rc))
                temp_lxx = []  # 用来存放列基本信息
                for xx in temp_bxx:
                    temp_lxx.append(xx)
                for cell in row.cells:
                    if rc != 1:
                        wb = cell.text.replace('\n', '')
                        temp_lxx.append(wb)
                # print("要插入的列数据是：", temp_lxx)  # 注释最长是1kb大致上三百字
                if rc > 1:
                    temp_altercolumncommentsql = "alter table sc_jc_gszj.cx_ypt_sjcc_{table_name} change column {column_name} comment '{column_comment}';".format(
                        table_name=temp_lxx[1], column_name=temp_lxx[3], column_comment=temp_lxx[4])
                    temp_altercolumncommentsql = temp_altercolumncommentsql.upper()
                    if re.findall('WBJH', temp_altercolumncommentsql):
                        print(temp_altercolumncommentsql)

def getData():
    str='表格ZS_LXGZ_ZSXXMXB的列清单'
    temp1=len(re.findall('的列清单',str))
    temp=re.findall('[A-Za-z_]+',str)[0]
    print(temp1)

if __name__ == '__main__':
    getData()

