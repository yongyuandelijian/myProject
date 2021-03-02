import docx
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

#document = Document()
# 设置一个空白样式
#style = document.styles['Normal']
# 设置西文字体
#style.font.name = 'Times New Roman'
## 设置中文字体
#style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

#定义 创建 超链接 函数
def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink
#定义 创建 书签 函数
def add_bookmark(paragraph, bookmark_text, bookmark_name):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: 创建一个段落
    :param bookmark_text: 定位到书签文本会被插入到文档中，
    :param tebookmark_name: 书签名称
    """
    run = paragraph.add_run()
    tag = run._r  # for reference the following also works: tag =  document.element.xpath('//w:r')[-1]
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), '0')
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(start)

    text = docx.oxml.OxmlElement('w:r')
    text.text = bookmark_text
    tag.append(text)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), '0')
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    tag.append(end)

document = docx.Document()
p = document.add_paragraph()
paragraph= document.add_paragraph()

document.add_page_break()

# 添加链接到本文档到书签链接 #temp1 #号后面跟书签名称
hyperlink = add_hyperlink(p, '#temp1', '书签', None, True)
#分页
document.add_page_break()

# 添加链接到url
hyperlink = add_hyperlink(p, 'www.baidu.com', '百度', None, True)

#可以将插入段落放在 add_paragraph('1') 也可以  将插入段落 放在 bookmark_text  书签名称bookmark_name
a=add_bookmark(paragraph= document.add_paragraph('1'), bookmadrk_text="",bookmark_name="temp1")
document.add_page_break()

b=add_bookmark(paragraph= document.add_paragraph('2'), bookmark_text="",bookmark_name="temp2")

document.add_page_break()
c=add_bookmark(paragraph= document.add_paragraph('3'), bookmark_text="",bookmark_name="temp3")
document.add_page_break()

document.save('test.docx')