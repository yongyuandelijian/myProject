# ecoding:utf-8
# 功能： 公共方法用作父类
# aaa  20210219

import os
import shutil
from docx.oxml.shared import OxmlElement  # 用于添加链接
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_LINE_SPACING,WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt,RGBColor
import win32com
from win32com.client import Dispatch

class Common():
    # 子类都必须要有这个参数
    def __init__(self, ny, root_dir, wordName):
        self.ny = ny
        self.root_dir = root_dir
        self.wdmc = os.path.join(root_dir, wordName)

    def moveoldfile(self):
        '''用于移动历史附件到old中，方便处理本次数据'''
        curdir = os.path.join(self.root_dir, "appendix")
        tardir = os.path.join(curdir, 'old')
        # 如果目标目录不存在进行创建
        if not os.path.exists(tardir):
            os.makedirs(tardir)
        # 获取目录中的文件进行移动
        curpath_files = os.listdir(curdir)
        # print(curpath_files)
        for filename in curpath_files:
            fullpath = os.path.join(curdir, filename)
            # print(fullpath)
            if os.path.isfile(fullpath):
                print("正在移动：", fullpath)
                temp = os.path.join(tardir, filename)
                if os.path.isfile(temp):
                    os.remove(temp)
                shutil.move(fullpath, tardir)

        print("目录处理完毕")

        # 创建一个书签函数，用来实现目录导航

    def add_bookmark(self, paragraph, bookmark_text, bookmark_name):
        '''
        功能：创建一个书签 在wps中不起作用
        参数：段落，书签文本（标记连接和段落的字符串），书签名称
        用法：
        '''
        run = paragraph.add_run()
        tag = run._r
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), bookmark_name)
        tag.append(start)

        text = OxmlElement('w:r')
        text.text = bookmark_text
        tag.append(text)

        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        end.set(qn('w:name'), bookmark_name)
        tag.append(end)

    def addTitle(self,doc,ztdx,nr,hjjbs=1.73,dljg=13,jb=3,fontname=u'宋体'):
        """添加一个标题"""
        p = doc.add_heading(level=jb)
        run = p.add_run(nr)
        run.font.name = fontname
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)
        jbzt = run.font
        jbzt.bold = True
        jbzt.color.rgb = RGBColor(0, 0, 0)
        jbzt.size = Pt(ztdx)
        pf = p.paragraph_format
        pf.space_before = Pt(dljg)
        pf.space_after = Pt(dljg)
        pf.line_spacing_rule=WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = hjjbs

    def addParagraph(self, nr, ztdx, doc):
        """添加一个标准的段落"""
        p = doc.add_paragraph()
        run = p.add_run(nr)
        zt = run.font
        zt.size = Pt(ztdx)
        pf = p.paragraph_format
        pf.first_line_indent = Inches(0.3)
        dljg=(ztdx+2) * 0.5
        pf.space_before = Pt(dljg)
        pf.space_after = Pt(dljg)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def addSpecialParagraph(self, nr, ztdx, doc,sp):
        """添加一个段落，将分隔符之前进行加粗，后面的保持不变,传入的时候注意分隔符的中英文区别"""
        p = doc.add_paragraph()
        nrlist=nr.split(sp,1)
        # 如果能分割，说明分隔符存在，否则就按普通处理，方法处理后会将分隔符省去，在后面拼接上
        # print(nrlist)
        if len(nrlist)==2:
            run1 = p.add_run(nrlist[0]+"：")
            zt1 = run1.font
            zt1.size = Pt(ztdx)
            zt1.bold=True
            run2 = p.add_run(nrlist[1])
            zt2 = run2.font
            zt2.size = Pt(ztdx)
            zt2.bold = False
        else:
            run1 = p.add_run(nr)
            zt1 = run1.font
            zt1.size = Pt(ztdx)
        pf = p.paragraph_format
        pf.first_line_indent = Inches(0.3)
        dljg = (ztdx + 2) * 0.5
        pf.space_before = Pt(dljg)
        pf.space_after = Pt(dljg)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def addSpecialParagraph2(self, nr, ztdx, doc,sp_before,sp_after):
        """添加一个段落，将两个分隔符之间的内容进行加粗，其他保持不变,传入的时候注意分隔符的中英文区别"""
        nrlist = nr.split(sp_after, 1)
        if len(nrlist) == 2:
            temp = nrlist[0]
            templist = temp.split(sp_before, 1)
            if len(templist) == 2:
                p = doc.add_paragraph()
                run0 = p.add_run(templist[0] + sp_before)
                zt0 = run0.font
                zt0.size = Pt(ztdx)
                zt0.bold = False
                run1 = p.add_run(templist[1] + sp_after)
                zt1 = run1.font
                zt1.size = Pt(ztdx)
                zt1.bold = True
                run2 = p.add_run(nrlist[1])
                zt2 = run2.font
                zt2.size = Pt(ztdx)
                zt2.bold = False
                pf = p.paragraph_format
                pf.first_line_indent = Inches(0.3)
                dljg = (ztdx + 2) * 0.5
                pf.space_before = Pt(dljg)
                pf.space_after = Pt(dljg)
                pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            else:
                self.addParagraph(nr,ztdx,doc)
        else:
            self.addParagraph(nr,ztdx,doc)


    def setTabBgColor(self, table, colnum, colorStr="A9A9A9"):
        """当然添加表格首行颜色肯定可以和添加表格合并，但是暂时没必要,在里面引用即可，这个对效率影响非常大，十几秒的代码增加这两行之后需要六十多秒"""
        shading_list = locals()
        for i in range(colnum):
            shading_list['shading_elm_' + str(i)] = parse_xml(
                r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
            table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_' + str(i)])

    def addWarnParagraph(self,doc,nr="此处对应的sql未查询到内容，请注意检查"):
        """添加一个警告段落，用于无法查询到对应数据的的时候提醒对应人员进行核实"""
        p = doc.add_paragraph()
        run = p.add_run(nr)
        zt = run.font
        zt.size = Pt(18)
        zt.color.rgb=RGBColor(255,0,0)
        pf = p.paragraph_format
        pf.first_line_indent = Inches(0.3)
        pf.space_before = Pt(36)
        pf.space_after = Pt(36)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def addTable(self, doc, data, table_style, tr_text, tr_width, ztdx):
        """添加一个标准的表格"""
        colnum = len(tr_text)
        table = doc.add_table(rows=1, cols=colnum, style=table_style)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 处理表头
        for ri in range(colnum):
            hdr_cells = table.rows[0].cells
            hdr_cells[ri].text = tr_text[ri]
            hdr_cells[ri].width = Inches(tr_width[ri])
            cp = table.cell(0, ri).paragraphs[0]
            cp.runs[0].font.bold = True
            cp.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells[ri].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # 如果有数据在处理数据行
        if data:
            for ri in range(len(data)):
                row = table.add_row()
                for ci in range(len(data[ri])):
                    if ci == 0 and data[ri][0] == data[ri - 1][0]:
                        # 第一列如果和上一个内容重复，则将表格合并,当前单元格不用再执行任何代码
                        table.cell(ri, 0).merge(table.cell(ri + 1, 0))
                        continue
                    rowcells = row.cells
                    # print(ri,ci,data[ri][ci])
                    rowcells[ci].text = str(data[ri][ci])
                    cp = rowcells[ci].paragraphs[0]
                    run = cp.runs[0]
                    font = run.font
                    font.size = Pt(ztdx)
                    pf = cp.paragraph_format
                    rowcells[
                        ci].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 这个和源码的例子是不同的，不要直接抄源码例子的过来，否则会报错
                    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    self.setTabBgColor(table, colnum)
                    if ci != 0:
                        pf.alignment = WD_TABLE_ALIGNMENT.LEFT
        else:
            self.addWarnParagraph(doc,"此处对应的sql未查询到内容，请注意检查")

    def jszs(self,data,xb):
        """功能：用来对二维列表中指定的列进行求和"""
        temp_num = 0
        if data:
            for a in data:
                if xb > len(a):
                    raise ValueError('参数下标不能超过迭代对象第二级的长度！！！')
                else:
                    temp_num += a[xb]
        return temp_num


    def update_pagenum(self, wordPath, ztdx, doc):
        """更新word页码情况，获取到了总共的，但是不能和每一页对应起来，只能先放着了"""
        # win32com操作word的一些基本设置
        w = win32com.client.Dispatch("Word.Application")
        w.Visible = 0
        w.DisplayAlerts = 0
        # 使用win32com获取页码情况
        win32doc = w.Documents.Open(wordPath)
        sumNum = w.ActiveDocument.ComputeStatistics(2)

        # 使用docx更新页脚情况
        footer = doc.sections[0].footer
        yjdl = footer.paragraphs[0]
        run = yjdl.add_run()
        temp = "共{sumNum}页".format(sumNum=sumNum)
        print(temp)
        run.add_text(temp)
        font = run.font
        font.size = Pt(ztdx)
        yjdl.style = doc.styles["Header"]
        yjgs = yjdl.paragraph_format
        yjgs.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        win32doc.SaveAs(wordPath)
        win32doc.Close()

    def wordtopdf(self, wordname):
        """将指定的word转为pdf"""
        word = Dispatch('Word.Application')
        doc = word.Documents.Open(wordname)
        doc.SaveAs(wordname.replace(".docx", ".pdf"), FileFormat="wdFormatPDF")  # 此方法必须安装word服务
        doc.Close()
        word.Quit()
