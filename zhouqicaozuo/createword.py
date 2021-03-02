# ecoding:utf-8
# 不加上面这个突然就有问题
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement  # 用于添加链接
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_CELL_VERTICAL_ALIGNMENT
# from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches,RGBColor,Pt,Cm
from functools import reduce
from collections import defaultdict
from pyecharts import options as opts
from pyecharts.render import make_snapshot
from pyecharts.charts import Map,Bar,Line,Pie
from pyecharts.globals import CurrentConfig   # 用于设置pyecharts读取本地js，而不会再去网上下载
from matplotlib import pyplot as plt
from wbnr.yb_text import get_ybwb
from sqlcommand.yb_sql import sqlstr
from sqlcommand.getsqldata import getsqldata
from snapshot_selenium import snapshot
import xlsxwriter
import calendar
import time
import re
import os
import shutil

class scword():
    """
    功能：用于生成组内月报，这个对于周期性的文档报告很有意义
    author:aaa
    date:20200527
    标准：大标题23 一级标题22 二级标题15 三级标题12 一般字体直接输出 行距1.5倍，段前5p，段后10p  中文:宋体 英文 Times New Roman
    实现思路：三级架构，准备好每个一级目录下的二级目录和二级所对应的内容对象，然后统一按照结构循环输出
    问题：1 Excel如何插入进来,2 目录如何生成,3页码如何获取 如果后续有方法了需要逐项解决，部分样式在wps中不起作用

    修改于 20210119
    第二版几乎所有结构全部变化，所以后续如果改版结构变化，则需要仔细看这里的逻辑#
    当前第二版，几乎没啥规律已经不是太适合使用循环了，每一章单独处理
    变化1：结构：第1，4章两级标题加内容 第二张四级标题加内容，第三章 三级标题加内容
    变化2：不再是之前的统一字体，需要仔细去看每一级别中包括页眉中使用字体的大小和类型（黑体/宋体），且每一级标题也会有不同，请注意查看
    变化3：存在一个小问题：文档模板要求每一级标题和上下级标题联动，这里不能自动编号，是手动写成的，请注意
    变化4：字体大小和类型都变化之后，那么需要每一级级别去查看，中文的字号和英文的磅需要去查看对照关系来设置
    注意：如果后续有新的图片，截图的时候，先要将正文页面放到最大，然后再去截图，否则容易导致客户放大后，图片中的文字失真
    """
    """
    当前版本存在的问题：
    1 重复内容表格合并-这个看情况吧，需要了在加
    2 excel不能自动插入- 需要手动插入到对应的位置
    3 不能自动获取页码-导致生成的目录没有页码，以及页脚的页码需要手动
    """
    def __init__(self,ny):
        self.ny=ny
        self.wdmc=r'../appendix/云平台运行维护及优化完善项目-数据集成月报-{ny}.docx'.format(ny=ny)

    def moveoldfile(self):
        '''用于移动历史附件到old中，方便处理本次数据'''
        curdir=os.path.abspath("../appendix")
        tardir=os.path.join(curdir,'old')
        # 如果目标目录不存在进行创建
        if not os.path.exists(tardir):
            os.makedirs(tardir)
        # 获取目录中的文件进行移动
        curpath_files=os.listdir(curdir)
        # print(curpath_files)
        for filename in curpath_files:
            fullpath=os.path.join(curdir,filename)
            # print(fullpath)
            if os.path.isfile(fullpath):
                print("正在移动：",fullpath)
                temp=os.path.join(tardir,filename)
                if os.path.isfile(temp):
                    os.remove(temp)
                shutil.move(fullpath,tardir)

        print("目录处理完毕")

    # 创建一个书签函数，用来实现目录导航
    def add_bookmark(self,paragraph,bookmark_text,bookmark_name):
        '''
        功能：创建一个书签 在wps中不起作用
        参数：段落，书签文本（标记连接和段落的字符串），书签名称
        用法：
        '''
        run=paragraph.add_run()
        tag=run._r
        start=OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'),'0')
        start.set(qn('w:name'),bookmark_name)
        tag.append(start)

        text=OxmlElement('w:r')
        text.text=bookmark_text
        tag.append(text)

        end=OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'),'0')
        end.set(qn('w:name'),bookmark_name)
        tag.append(end)

    def addParagraph(self,nr,ztdx,doc):
        """添加一个标准的段落"""
        p = doc.add_paragraph()
        run = p.add_run(nr)
        zt = run.font
        zt.size = Pt(ztdx)
        pf = p.paragraph_format
        pf.first_line_indent = Inches(0.3)
        pf.space_before = Pt(ztdx * 0.5)
        pf.space_after = Pt(ztdx * 0.5)
        pf.line_spacing = Pt(ztdx * 1.5)

    def setTabBgColor(self,table,colnum,colorStr="A9A9A9"):
        """当然添加表格首行颜色肯定可以和添加表格合并，但是暂时没必要,在里面引用即可，这个对效率影响非常大，十几秒的代码增加这两行之后需要六十多秒"""
        shading_list=locals()
        for i in range(colnum):
            shading_list['shading_elm_'+str(i)]=parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'),bgColor=colorStr))
            table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_'+str(i)])

    def addTable(self,doc,data,table_style,tr_text,tr_width,ztdx):
        """添加一个标准的表格"""
        colnum=len(data[0])
        table = doc.add_table(rows=1, cols=colnum, style=table_style)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 处理表头
        for ri in range(colnum):
            hdr_cells = table.rows[0].cells
            hdr_cells[ri].text = tr_text[ri]
            hdr_cells[ri].width = Inches(tr_width[ri])
            cp = table.cell(0, ri).paragraphs[0]
            cp.runs[0].font.bold = True
            cp.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells[ri].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        # 处理数据行
        for ri in range(len(data)):
            row = table.add_row()
            for ci in range(len(data[ri])):
                rowcells = row.cells
                # print(ri,ci,data[ri][ci])
                rowcells[ci].text = str(data[ri][ci])
                cp = rowcells[ci].paragraphs[0]
                run=cp.runs[0]
                font = run.font
                font.size = Pt(ztdx)
                pf = cp.paragraph_format
                rowcells[ci].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if ci!=0:
                    pf.alignment = WD_TABLE_ALIGNMENT.LEFT

                pf.line_spacing = Pt(ztdx * 1.5)
                # table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                self.setTabBgColor(table,colnum)

    # 添加一个左右分割开的页眉样式
    # from docx.enum.style import WD_STYLE_TYPE
    # styles=doc.styles
    # style=styles.add_style("myheadstyle",WD_STYLE_TYPE)
    # style.base_style=styles["Normal"]
    # tab_stops=style.paragraph_format.tab_stops


    doc = Document()  # 生成一个word对象
    # 设置pyecharts读取的路径
    js_path = r"../js/"
    CurrentConfig.ONLINE_HOST = js_path

    def scword(self):
        '''生成文档'''
        self.moveoldfile()  # 处理附件目录
        # sql查询数据
        getdata_dx = getsqldata()  # 获取sql的类对象
        sqltext = sqlstr(self.ny)  # 获取sql文本类对象
        # 公共部分
        doc=self.doc
        style = doc.styles['Normal']
        style.font.name = u'Times New Roman'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        style.font.color.rgb=RGBColor(0,0,0)
        nf=int(self.ny[:4])
        yf=int(self.ny[-2:])
        start='1'
        end = calendar.monthrange(nf,yf)[1]
        table_style = "Table Grid"
        yczs = 0  # 临时做一些统计数留作下一模块使用
        # 各个级别的字体对照大小
        ztdx_jb1 = 22
        ztdx_jb2 = 16
        ztdx_jb3 = 15
        ztdx_jb4 = 14
        ztdx_ym=9
        ztdx_nr = 12
        ztdx_5h=10.5

        # 设置matplotlib 中文显示问题
        plt.rcParams['font.sans-serif'] = ['KaiTi']  # 制定默认字体
        plt.rcParams['axes.unicode_minus'] = False   # 解决保存图像符号显示方块的问题
        ################################## 添加文档页眉 ##################################
        header=doc.sections[0].header
        ymdl=header.paragraphs[0]
        run=ymdl.add_run()
        run.add_picture(r'../img/img_ym01.png',width=Cm(2.49),height=Cm(0.9))
        ymnr='\t'+' '*45+'2019年云平台运行维护及优化完善项目-数据层运维统计分析月报'
        run.add_text(ymnr)
        font=run.font
        font.size=Pt(ztdx_ym)
        run.underline = True
        ymdl.style=doc.styles["Header"]
        ymgs=ymdl.paragraph_format
        ymgs.tab_stops.add_tab_stop(Cm(2),WD_TABLE_ALIGNMENT.RIGHT,WD_TAB_LEADER.DOTS)  # 从中间隔开，另一边右对齐SPACES 没有作用，还需要研究
        ymgs.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER

        ################################## 封面 ##################################
        paragraph = doc.add_paragraph()
        dbt = paragraph.add_run('云平台运行维护及优化完善项目\n\n数据集成月报')
        # 添加run对象，参数为text=None和style=None,
        dbt.bold=True
        dbt.font.name=u'黑体'
        dbt._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        # 字体对象
        btzt = dbt.font
        btzt.size=Pt(ztdx_jb1)  # 设置标题字体大小
        btzt.bold=True
        btzt.color.rgb = RGBColor(0, 0, 0)
        paragraph_format=paragraph.paragraph_format
        paragraph_format.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format.space_after=Pt(ztdx_jb1*6)
        paragraph_format.space_before=Pt(ztdx_jb1*3)
        paragraph_format.line_spacing=Pt(ztdx_jb1*1.5)

        # 封面最底下的一行宋体
        paragraph=doc.add_paragraph()
        dbt2=paragraph.add_run('中国软件与技术服务股份有限公司')
        dbt2.font.name=u'宋体'
        dbt2._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        btzt2 = dbt2.font
        btzt2.size = Pt(ztdx_jb2)
        paragraph_format=paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_page_break()

        ################################## 目录部分 ##################################

        paragraph = doc.add_paragraph()
        xdbt = paragraph.add_run("目 录")
        xdbtzt = xdbt.font
        xdbtzt.name = u'黑体'
        xdbt._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        xdbtzt.size = Pt(22)
        xdbtzt.bold=True
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # # 段落前后的空行数量
        paragraph_format.space_after=Pt(24)
        paragraph_format.space_before = Pt(16)

        # 获取到固定内容和动态内容的查询sql对象
        nr21,all_text=get_ybwb.hqnr(get_ybwb())   # 获取到结构内容和直属内容
        # sql = sqlstr(self.ny)
        # print("正在生成的是{ny}的数据，请注意核实".format(ny=sql.ny))

        # 一级目录加粗，二级目录直接输出,三级目录斜体，使用默认宋体
        for ml_jb1 in all_text.keys():
            # print(ml_jb1)
            p1=doc.add_paragraph(ml_jb1)
            dl_jb1=p1.add_run()  # 都是用的这个方法，但是\t只在这里一层这里有bug导致\t之后的内容未显示出来，移动到上面后正常，需要注意
            # self.add_bookmark(p1,ml_jb1,ml_jb1)
            dlzt_jb1=dl_jb1.font
            dlzt_jb1.bold=True
            dlzt_jb1.size=Pt(10)
            # 如果是2和3要单独处理,2虽然有4级但是保持模板一致只打印三级
            if ml_jb1[:ml_jb1.index("\t")]=="第2章" or ml_jb1[:ml_jb1.index("\t")]=="第3章":
                for ml_jb2 in all_text.get(ml_jb1):
                    for jb2key in ml_jb2.keys():
                        p2 = doc.add_paragraph()  # 直接加入到段落里，会造成空格制表符等缺失，这个看自己是否需要
                        r2=p2.add_run(jb2key)
                        # self.add_bookmark(p2,jb2key,jb2key)
                        pf2 = p2.paragraph_format
                        pf2.first_line_indent = Inches(0.3)
                        if jb2key[:jb2key.index("\t")]=='2.2' or jb2key[:jb2key.index("\t")]=='2.3':
                            continue
                        for ml_jb3 in ml_jb2.get(jb2key):
                            for jb3key in ml_jb3.keys():
                                p3 = doc.add_paragraph()
                                r3=p3.add_run(jb3key)
                                # self.add_bookmark(p3,jb3key,jb3key)
                                r3.Italic=True
                                pf3 = p3.paragraph_format
                                pf3.first_line_indent = Inches(0.6)  # 缩进双倍

            else:
                for xl in all_text.get(ml_jb1):
                    for xj in xl.keys():
                        paragraph=doc.add_paragraph(xj)
                        paragraph_format=paragraph.paragraph_format
                        paragraph_format.first_line_indent=Inches(0.3)
        doc.add_page_break()

        ################################## 添加正文，这里是打印所有目录内容 ##################################
        # 第二版几乎所有结构全部变化，所以后续如果改版，则需要仔细看这里的逻辑#
        # 当前第二版，几乎没啥规律，1，4章两级标题加内容 第二张四级标题加内容，第三章 三级标题加内容，
        # 内容字体规律：第一级标题宋体2号 22，第二级标题黑体3号 16，第三级宋体小3 15，第四级，宋体4号 14，内容宋体小4 12

        for jb1 in all_text.keys():
            dldl = doc.add_heading(level=1)  # 添加标题
            pf=dldl.paragraph_format
            pf.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
            pf.space_before=Pt(17)
            pf.space_after=Pt(17)
            pf.line_spacing=Pt(ztdx_jb1*2.41)
            # run=dldl.add_run(jb1[:jb1.rindex("\t")])
            run = dldl.add_run(jb1)
            run.font.name = u'宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            dlbtzt=run.font
            dlbtzt.bold = True
            dlbtzt.size = Pt(ztdx_jb1)
            # 设置颜色，这两种都可以
            dlbtzt.color.rgb=RGBColor(0,0,0)  # WD_COLOR_INDEX.AUTO
            # dlbtzt.color.theme_color=WD_COLOR_INDEX.BLACK

            # 处理第二级
            for jb2 in all_text.get(jb1):
                for jb2key in jb2.keys():
                    p2 = doc.add_heading(level=2)
                    run2 = p2.add_run(jb2key)
                    run2.font.name = u'黑体'
                    run2._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
                    jb2zt = run2.font
                    jb2zt.bold = True
                    jb2zt.size = Pt(ztdx_jb2)
                    jb2zt.color.rgb=RGBColor(0,0,0)
                    pf2 = p2.paragraph_format
                    pf2.space_before=Pt(13)
                    pf2.space_after=Pt(ztdx_jb2*0.5)
                    pf2.line_spacing=Pt(ztdx_jb2*1.73)
                    if jb2key[:jb2key.index("\t")] == "2.1":
                        # 处理直属内容
                        doc.add_picture(r"..\img\img_dl21.bmp", width=Inches(6.25))  # 这个数字必须指定
                        for nr in nr21:
                            self.addParagraph(nr, ztdx_nr, doc)
                        # 处理比其他章节多的层级
                        # 处理第三级
                        for jb3 in jb2.get(jb2key):
                            for jb3key in jb3.keys():
                                p3=doc.add_heading(level=3)
                                run3=p3.add_run(jb3key)
                                run3.font.name = u'宋体'
                                run3._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                                jb3zt = run3.font
                                jb3zt.bold=True
                                jb3zt.color.rgb=RGBColor(0,0,0)
                                jb3zt.size=Pt(ztdx_jb3)
                                pf3=p3.paragraph_format
                                pf3.space_before=Pt(13)
                                pf3.space_after=Pt(13)
                                pf3.line_spacing=Pt(ztdx_jb3*1.73)
                                # 处理第四级
                                for jb4 in jb3.get(jb3key):
                                    for jb4key in jb4.keys():
                                        p4=doc.add_heading(level=4)
                                        run4=p4.add_run(jb4key)
                                        run4.font.name = u'宋体'
                                        run4._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                                        jb4zt=run4.font
                                        jb4zt.bold=True
                                        jb4zt.italic=False
                                        jb4zt.color.rgb=RGBColor(0,0,0)
                                        jb4zt.size=Pt(ztdx_jb4)
                                        pf4=p4.paragraph_format
                                        pf4.space_before=Pt(ztdx_jb4)
                                        pf4.space_after=Pt(ztdx_jb4)
                                        pf4.line_spacing=Pt(ztdx_jb4*1.57)
                                    for nr in jb4.values():
                                        self.addParagraph(nr,ztdx_nr,doc)
                    elif jb2key[:jb2key.index("\t")] == "2.2":
                        # 插入表2.2 获取到的数据元组套元组-问题：列宽只支持office
                        data=getsqldata.getdata(getdata_dx,sqltext.get_b22_jcqk())
                        tr_text = ["分类", "来源系统", "表数量", "数据量（万）", "存储（GB）", "集成频率", "集成方式"]
                        tr_width=[1,2,1,3,2,0.75,0.75]
                        self.addTable(doc,data,table_style,tr_text,tr_width,ztdx_5h)
                        # 正常打印段落
                        for nr in jb2.get(jb2key):
                            self.addParagraph(nr,ztdx_nr,doc)
                    elif jb2key[:jb2key.index("\t")] == "2.3":
                        temp=0
                        for nr in jb2.get(jb2key):
                            if temp==1:
                                doc.add_picture(r"..\img\img_dl23.bmp",width=Inches(6))
                            self.addParagraph(nr,ztdx_nr,doc)
                            temp = temp + 1
                    elif jb2key[:jb2key.index("\t")] == "3.1" or jb2key[:jb2key.index("\t")] == "3.2":
                        for jb3 in jb2.get(jb2key):
                            for jb3key in jb3.keys():
                                p3 = doc.add_heading(level=3)
                                run3 = p3.add_run(jb3key)
                                run3.font.name = u'宋体'
                                run3._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
                                jb3zt = run3.font
                                jb3zt.size = Pt(ztdx_jb3)
                                jb3zt.color.rgb=RGBColor(0,0,0)
                                pf3 = p3.paragraph_format
                                pf3.space_before = Pt(13)
                                pf3.space_after = Pt(13)
                                pf3.line_spacing = Pt(ztdx_jb3 * 1.73)
                                # 输出此段需要单独处理的内容

                                if jb3key[:jb3key.index("\t")]=="3.1.1":
                                    nr=jb3.get(jb3key)[0]  # 此段落目前只有一个内容所以直接取出来
                                    self.addParagraph(nr,ztdx_nr,doc)
                                    # 处理动态内容
                                    data=getsqldata.getdata(getdata_dx,sqltext.get_b311_ycjc())
                                    tr_text=['日期','异常进程数','异常进程单位']
                                    tr_width=[2,3,5]
                                    self.addTable(doc,data,table_style,tr_text,tr_width,ztdx_5h)
                                    # 统计文字
                                    valid_text=""
                                    temp_num=0
                                    for hang in data:
                                        if len(hang[2])>0:
                                            temp="{ri}号有{cs}次异常,".format(ri=hang[0][-2:],cs=hang[1])
                                            valid_text=valid_text+temp
                                            yczs=yczs+int(hang[1])
                                        else:
                                            temp_num=temp_num+1
                                    nr = "分日期统计可以看出{yf}月份从{start}号到{end}号每天的任务运行失败情况，其中异常情况依次是:{valid_text}{yf}月共有{num}天没有失败的任务。".format(yf=yf,start=start,end=end,valid_text=valid_text,num=temp_num)
                                    self.addParagraph(nr,ztdx_nr,doc)
                                elif jb3key[:jb3key.index("\t")] == "3.1.2":
                                    nr = jb3.get(jb3key)[0]  # 此段落目前只有一个内容所以直接取出来
                                    self.addParagraph(nr,ztdx_nr,doc)
                                    # 处理动态内容,后期这个表格这部分固定内容可以抽象出去的
                                    tr_text = ['序号', 'OGG进程异常原因', '次数']
                                    row1=['1','版本发布导致：未按版本发布规范进行操作导致进程异常',yczs]
                                    row2=['2','系统故障：分发库机器检修及扩容','0']
                                    data=[row1,row2]
                                    tr_width = [2, 5, 2]
                                    self.addTable(doc,data,table_style,tr_text,tr_width,ztdx_5h)
                                    nr="如上所示：整个{yf}月份分发库OGG进程异常共有{yczs}次，通过对异常分析，得出产生问题的主要原因是分发库版本升级后未按规定步骤刷新定义文件导致解析进程异常。".format(yf=yf,yczs=yczs)
                                    self.addParagraph(nr,ztdx_nr,doc)
                                elif jb3key[:jb3key.index("\t")] == "3.2.1":
                                    # 添加固定的
                                    for nr in jb3.get(jb3key):
                                        self.addParagraph(nr,ztdx_nr,doc)
                                    # 添加动态的
                                    data=getsqldata.getdata(getdata_dx,sqltext.get_zj321_xzrwqk())[0]
                                    nr = "从{yf}月{start}号到{yf}月{end}号总共新增了{rwzs}个任务（基础层新增{bsl}张表，部署任务{dcrws}个；镜像层新增{bsl}张表，部署任务{dcrws}个任务）。".format(yf=yf,start=start,end=end,rwzs=data[0],bsl=data[2],dcrws=data[1])
                                    self.addParagraph(nr,ztdx_nr,doc)

                                    data=getsqldata.getdata(getdata_dx,sqltext.get_b321_rwyc_frq())
                                    day=[]
                                    ycs=[]
                                    for mx in data:
                                        day.append(mx[0])
                                        ycs.append(mx[1])
                                    bt="图3.2.1 {ny}分日任务异常进程情况统计图".format(ny=ny)
                                    zxt321 = Line(init_opts=opts.InitOpts(width="800px", height="400px"))
                                    (

                                        zxt321.add_xaxis(xaxis_data=day)
                                            .add_yaxis(
                                            series_name="",
                                            y_axis=ycs,
                                            linestyle_opts=opts.LineStyleOpts(color="green", width=4),
                                            markpoint_opts=opts.MarkPointOpts(
                                                data=[
                                                    opts.MarkPointItem(type_="max", name="最大值"),
                                                    opts.MarkPointItem(type_="min", name="最小值"),
                                                ]
                                            ),
                                            markline_opts=opts.MarkLineOpts(
                                                data=[opts.MarkLineItem(type_="average", name="平均值")]
                                            ),
                                        )
                                            .set_global_opts(
                                            title_opts=opts.TitleOpts(title=bt, subtitle=self.ny),
                                            xaxis_opts=opts.AxisOpts(name='(当月日期)'),
                                            yaxis_opts=opts.AxisOpts(name='（异常次数）')
                                        )
                                    )
                                    make_snapshot(snapshot, zxt321.render(), r"..\img\img_b321_rwyc_frq.png")
                                    self.doc.add_picture(r"..\img\img_b321_rwyc_frq.png", width=Inches(6.25))

                                    # 总结
                                    zjlist=sorted(data,key=lambda x:(x[1],x[0]),reverse=True)  # 利用sorted自动解包后进行排序，sorted默认按照第一个元素进行排序，所以要将排序的元素利用lambda放到前面来
                                    nr="分日期统计可以看出{yf}月份从{start}号到{end}号每天的任务运行失败情况，其中失败任务数排在前三位依次是:".format(yf=yf,start=start,end=end)
                                    temp_num=0
                                    for row in zjlist:
                                        if row[1]==0:
                                            temp="{yf}月共有{ts}天没有失败的任务。".format(yf=yf,ts=end-temp_num)
                                            break
                                        if temp_num<3:
                                            temp="{yf}月{rq}号有{ycs}次异常,".format(yf=yf,rq=row[0][-2:],ycs=row[1])
                                            nr=nr+temp
                                        temp_num+=1
                                    # 任务分单位异常情况表
                                    nr="{nf}年{yf}月任务分单位异常情况：".format(nf=nf,yf=yf)
                                    p = doc.add_paragraph()
                                    run = p.add_run(nr)
                                    zt = run.font
                                    zt.size = Pt(12)
                                    pf = p.paragraph_format
                                    pf.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
                                    pf.space_before = Pt(ztdx_nr * 0.5)
                                    pf.space_after = Pt(ztdx_nr * 0.5)
                                    pf.line_spacing = Pt(ztdx_nr * 1.5)

                                    data=getsqldata.getdata(getdata_dx,sqltext.get_b321_rwyc_fdw())
                                    tr_text=['序号','所属月份','单位名称','异常次数']
                                    tr_width=[1,3,3,3]
                                    self.addTable(doc,data,table_style,tr_text,tr_width,ztdx_5h)

                                    zjlist=sorted(data,key=lambda x:(x[3],x[0],x[1],x[2]),reverse=True)
                                    dws=0
                                    zcs=0
                                    temp=""
                                    dd=defaultdict(lambda :0)
                                    ditu_list = []
                                    for row in zjlist:
                                        if row[3]==0:
                                            break
                                        else:
                                            zcs+=row[3]
                                            temp+=row[2]+','
                                            dws += 1
                                            dd[row[2]]=row[3]
                                            ditu_list.append([row[2],row[3]])

                                    nr="{ny}整月集成问题共出现{zcs}次，共涉及{dws}家单位，分别是：{temp}".format(ny=ny,zcs=zcs,dws=dws,temp=temp)
                                    self.addParagraph(nr,ztdx_nr,doc)
                                    nr="区域地图显示如下："
                                    self.addParagraph(nr,ztdx_nr,doc)

                                    m1 = Map()
                                    m1.add(series_name="异常情况", data_pair=ditu_list, maptype="china",
                                           name_map={"key": "value"},
                                           is_map_symbol_show=True)
                                    m1.set_global_opts(
                                        title_opts=opts.TitleOpts(title="图3.2.1 任务异常情况分单位", subtitle=self.ny),
                                        # 设置标题和副标题在左侧
                                        visualmap_opts=opts.VisualMapOpts(is_piecewise=True, max_=max(dd.values()))
                                        # 是否分段,以及分段的最大值
                                    )
                                    make_snapshot(snapshot, m1.render(), r'..\img\img_nr321_rwycfdw.png')
                                    self.doc.add_picture(r'..\img\img_nr321_rwycfdw.png', width=Inches(6.25))

                                elif jb3key[:jb3key.index("\t")] == "3.2.2":
                                    # 添加固定的
                                    for nr in jb3.get(jb3key):
                                        self.addParagraph(nr, ztdx_nr, doc)
                                    nr = "\n任务主要异常原因分析"
                                    p = doc.add_paragraph()
                                    run = p.add_run(nr)
                                    zt = run.font
                                    run.blod=True
                                    zt.size = Pt(12)
                                    pf = p.paragraph_format
                                    pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                    pf.space_before = Pt(ztdx_nr * 0.5)
                                    pf.space_after = Pt(ztdx_nr * 0.5)
                                    pf.line_spacing = Pt(ztdx_nr * 1.5)
                                    # 如果超过5种原因，只获取前4，剩余都归到其他
                                    data=getsqldata.getdata(getdata_dx,sqltext.get_b322_rwyc_fyy())
                                    yccs = []
                                    ycyy = []
                                    explode = []
                                    temp_num=0
                                    qtsl=0
                                    for mx in data:
                                        if temp_num<5:
                                            yccs.append(mx[1])
                                            ycyy.append(mx[0])
                                            explode.append(0)
                                        else:
                                            qtsl+=mx[1]
                                        temp_num+=1

                                    if qtsl > 0:
                                        yccs.append(qtsl)
                                        ycyy.append("其他原因")
                                        explode.append(0)
                                    # 如果有数据增加异常原因饼图
                                    if len(yccs) > 0:
                                        explode[0] = 0.1  # 改变自己需要的分离缝隙  各个值所属模块突出来的缝隙大小
                                        fig1, ax1 = plt.subplots()
                                        ax1.pie(yccs, explode=explode, labels=ycyy, autopct='%1.2f%%',
                                                shadow=True, startangle=90)  # 分别是数据对象，爆裂值，名称列表，自动百分比，是否阴影，起始角度
                                        ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
                                        plt.savefig(r"..\img\img_t322_rwyc_fyy.png")
                                        self.doc.add_picture(r"..\img\img_t322_rwyc_fyy.png", width=Inches(6))

                                    # 添加总结概括的一段
                                    valid_text=reduce(lambda a,b:a+","+b,ycyy[:3])
                                    nr="如上图所示，{yf}月集成问题共出现{zcs}次，主要问题分类是：{valid_text}。导致失败的原因主要有三大类，一是分发库问题，主要是连接异常，版本不一致，脏数据（最大比重的就是脏数据问题）；二个是云平台问题，主要是实例调度问题、任务卡住、资源不足导致任务失败的问题，第三个主要是frp转发端口故障。从统计数量来看{maxyy}问题导致的失败任务是最多的,发生了{maxcs}次，这个问题需要各地运维人员严格按照规范来操作，也是后期需要解决的问题。".format(
                                        yf=yf,zcs=sum(yccs),valid_text=valid_text,maxyy=data[0][0],maxcs=data[0][1]
                                    )
                                    self.addParagraph(nr,ztdx_nr,doc)
                                elif jb3key[:jb3key.index("\t")] == "3.2.3":
                                    data=getsqldata.getdata(getdata_dx,sqltext.get_b323_sbrwydcl())
                                    tr_text=['序号','失败类别','解决方法']
                                    tr_width=[1,5,4]
                                    self.addTable(doc,data,table_style,tr_text,tr_width,ztdx_5h)
                                elif jb3key[:jb3key.index("\t")] == "3.2.4":
                                    temp_num=1
                                    for gdnr in jb3.get(jb3key):
                                        if temp_num==2:
                                            temp = "以下选取5张金三核心表对{yf}月24号至{yf}月28号5天数据抽取情况进行数据量分析比对，详情见附件：《{ny}数据抽样比对结果.xlsx》".format(yf=yf,ny=ny)
                                            self.addParagraph(temp,ztdx_nr,doc)
                                            # 生成excel
                                            list_tablename = ['HX_DJ_DJ_NSRXX', 'HX_DJ_DJ_NSRXX_KZ',
                                                              'HX_RD_RD_SFZRDXXB', 'HX_SB_SB_SBXX', 'HX_ZS_ZS_YJSF']
                                            workbook = xlsxwriter.Workbook(
                                                r"../appendix/{ny}数据抽样比对结果.xlsx".format(ny=self.ny))
                                            for tablename in list_tablename:
                                                worksheet = workbook.add_worksheet(tablename)
                                                # 声明一个加粗的样式用来增加给表头，颜色是16进制数
                                                bt = workbook.add_format(
                                                    {'bold': True, 'bg_color': '9BC2E6', 'align': 'centre',
                                                     'valign': 'vcentre',
                                                     'border': 1})
                                                nr = workbook.add_format(
                                                    {'align': 'centre', 'valign': 'vcentre',
                                                     'border': 1})
                                                # czjz=workbook.add_format({'align':'vcenter'})
                                                worksheet.merge_range(first_row=0, last_row=1, first_col=0, last_col=0,
                                                                      data="单位", cell_format=bt)
                                                worksheet.merge_range(first_row=0, last_row=0, first_col=1, last_col=4,
                                                                      data="{nf}年{yf}月24日".format(nf=nf, yf=yf),
                                                                      cell_format=bt)
                                                worksheet.merge_range(first_row=0, last_row=0, first_col=5, last_col=8,
                                                                      data="{nf}年{yf}月25日".format(nf=nf, yf=yf),
                                                                      cell_format=bt)
                                                worksheet.merge_range(first_row=0, last_row=0, first_col=9, last_col=12,
                                                                      data="{nf}年{yf}月26日".format(nf=nf, yf=yf),
                                                                      cell_format=bt)
                                                worksheet.merge_range(first_row=0, last_row=0, first_col=13,
                                                                      last_col=16,
                                                                      data="{nf}年{yf}月27日".format(nf=nf, yf=yf),
                                                                      cell_format=bt)
                                                worksheet.merge_range(first_row=0, last_row=0, first_col=17,
                                                                      last_col=20,
                                                                      data="{nf}年{yf}月28日".format(nf=nf, yf=yf),
                                                                      cell_format=bt)
                                                worksheet.write(1, 1, "分发库", bt)
                                                worksheet.write(1, 2, "基础层", bt)
                                                worksheet.write(1, 3, "差异", bt)
                                                worksheet.write(1, 4, "差异率", bt)
                                                worksheet.write(1, 5, "分发库", bt)
                                                worksheet.write(1, 6, "基础层", bt)
                                                worksheet.write(1, 7, "差异", bt)
                                                worksheet.write(1, 8, "差异率", bt)
                                                worksheet.write(1, 9, "分发库", bt)
                                                worksheet.write(1, 10, "基础层", bt)
                                                worksheet.write(1, 11, "差异", bt)
                                                worksheet.write(1, 12, "差异率", bt)
                                                worksheet.write(1, 13, "分发库", bt)
                                                worksheet.write(1, 14, "基础层", bt)
                                                worksheet.write(1, 15, "差异", bt)
                                                worksheet.write(1, 16, "差异率", bt)
                                                worksheet.write(1, 17, "分发库", bt)
                                                worksheet.write(1, 18, "基础层", bt)
                                                worksheet.write(1, 19, "差异", bt)
                                                worksheet.write(1, 20, "差异率", bt)
                                                # 获取数据进行添加
                                                data = getsqldata.getdata(getdata_dx, sqltext.get_b324_sjyzx(tablename))
                                                for row_num in range(len(data)):
                                                    for col_num in range(len(data[0])):
                                                        worksheet.write(row_num + 2, col_num, data[row_num][col_num],nr)
                                            workbook.close()
                                        if temp_num==3:
                                            temp = "另本月运维和数据差异的处理记录，详情见附件：《{ny}数据差异运维记录.xlsx》。".format(ny=ny)
                                            self.addParagraph(temp,ztdx_nr,doc)
                                            # 生成 4.5 excel 2 数据差异运维记录
                                            workbook = xlsxwriter.Workbook(
                                                "../appendix/{ny}数据差异运维记录.xlsx".format(ny=self.ny))
                                            worksheet = workbook.add_worksheet("运维记录")
                                            # 声明一个加粗的样式用来增加给表头，颜色是16进制数
                                            bt = workbook.add_format(
                                                {'bold': True, 'bg_color': '9BC2E6', 'align': 'centre',
                                                 'valign': 'vcentre', 'border': 1})
                                            nr = workbook.add_format(
                                                {'align': 'left', 'valign': 'vcentre',
                                                 'border': 1})
                                            worksheet.write(0, 0, "项目名称", bt)
                                            worksheet.write(0, 1, "业务日期", bt)
                                            worksheet.write(0, 2, "任务名称", bt)
                                            worksheet.write(0, 3, "问题类型", bt)
                                            worksheet.write(0, 4, "处理过程", bt)
                                            worksheet.write(0, 5, "错误原因", bt)
                                            worksheet.write(0, 6, "影响范围", bt)
                                            worksheet.write(0, 7, "发现日期", bt)
                                            worksheet.write(0, 8, "处理日期", bt)
                                            # 设置列宽
                                            worksheet.set_column(0, 1, 15)
                                            worksheet.set_column(2, 8, 30)
                                            data=getsqldata.getdata(getdata_dx,sqltext.get_b324_sjcyyw())
                                            for row_num in range(len(data)):
                                                for col_num in range(len(data[0])):
                                                    worksheet.write(row_num + 1, col_num, data[row_num][col_num], nr)
                                            workbook.close()
                                        self.addParagraph(gdnr,ztdx_nr,doc)
                                        temp_num+=1
                                elif jb3key[:jb3key.index("\t")] == "3.2.5":
                                    data = getsqldata.getdata(getdata_dx, sqltext.get_b325_sjgxqk())
                                    nr="目前云平台基础层配置调度任务的数据源共有49个，包含日调度、周调度、月调度三种更新频率，共计6803张表。根据本月监控源端数据变化情况统计，其中按调度更新的表有3093张，未更新的表有3710张。各源端系统详细情况如下表所示:"
                                    self.addParagraph(nr,ztdx_nr,doc)
                                    tr_width=[1,4,1,2,2]
                                    tr_text=['序号','数据来源','总量','更新表数据量','未更新表数量']
                                    self.addTable(doc,data,table_style,tr_text,tr_width,ztdx_5h)
                                    for nr in jb3.get(jb3key):
                                        self.addParagraph(nr,ztdx_nr,doc)
                                else:
                                    for nr in jb3.get(jb3key):
                                        self.addParagraph(nr,ztdx_nr,doc)
                    # 处理两级标题后正常内容的
                    else:
                        for nr in jb2.get(jb2key):
                            p = doc.add_paragraph()
                            run=p.add_run(nr)
                            zt=run.font
                            zt.size=Pt(12)
                            pf = p.paragraph_format
                            pf.first_line_indent = Inches(0.3)
                            pf.space_before = Pt(ztdx_nr * 0.5)
                            pf.space_after = Pt(ztdx_nr * 0.5)
                            pf.line_spacing = Pt(ztdx_nr * 1.5)

            if jb1[:jb1.index("\t")]!="第4章":
                doc.add_page_break()  # 除最后一章，每一章完成后换页
        doc.save(self.wdmc)

if __name__ == '__main__':
    starttime=time.time()
    ny='202101'
    if re.match('[2][0][0-9]{2}',ny[:4]) or re.match('[0,1][0-9]',ny[4:]):
        scword(ny).scword()
    else:
        print("年月输入异常，请处理后再执行脚本")
    print("生成完毕,消耗时间：",round(time.time()-starttime,2))