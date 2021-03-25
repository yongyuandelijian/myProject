# ecoding:utf-8
from docx import Document
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, RGBColor, Pt, Cm
from functools import reduce
from collections import defaultdict
from wbnr.yb_text_zlz import get_ybwb
from sqlcommand.yb_sql_zlz import sqlstr
from sqlcommand.getsqldata import getsqldata
from zhouqicaozuo.common import Common
import calendar
import os
import time


class scword(Common):
    """
    功能：按照给定的模板生成治理组的月报
    author:aaa
    date:20210219
    标准：大标题22 一级标题22 二级标题15 三级标题12 一般字体直接输出 行距1.5倍，段前5p，段后10p  中文:宋体 英文 Times New Roman
    实现思路：三级架构，准备好每个一级目录下的二级目录和二级所对应的内容对象，然后统一按照结构循环输出

    当前版本存在的问题：
    1 不能自动获取页码-所以不生成目录以及页脚，目录和页脚手动操作
    """
    # 公共变量
    doc = Document()  # 生成一个word对象

    # 重写公共方法内的表格方法，这里的样式上有一些差异，标题的字体是宋体小四，加粗，背景色是蓝色，内容是宋体5号 所有内容都是居中，当然了如果还是想要公共，还是可以传入更多参数，只是因为不愿修改前面的所以重写一下
    def addTable(self, doc, data, table_style, tr_text, tr_width, ztdx):
        """添加一个标准的表格"""
        colnum = len(tr_text)
        table = doc.add_table(rows=1, cols=colnum, style=table_style)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 处理表头
        for ri in range(colnum):
            hdr_cells = table.rows[0].cells
            hdr_cells[ri].text = tr_text[ri]
            hdr_cells[ri].width = Inches(tr_width[ri])
            cp = table.cell(0, ri).paragraphs[0]
            cp.runs[0].font.bold = True
            cp.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells[ri].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            self.setTabBgColor(table, colnum, "B4C6E7")
        # 如果有数据在处理数据行
        if data:
            for ri in range(len(data)):
                row = table.add_row()
                for ci in range(len(data[ri])):
                    rowcells = row.cells
                    nr = data[ri][ci]
                    if type(nr) == float:
                        nr = int(nr)
                    rowcells[ci].text = str(nr)
                    cp = rowcells[ci].paragraphs[0]
                    run = cp.runs[0]
                    font = run.font
                    font.size = Pt(ztdx)
                    pf = cp.paragraph_format
                    rowcells[ci].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 这个和源码的例子是不同的，不要直接抄源码例子的过来，否则会报错
                    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    pf.line_spacing = 1.5
                    pf.alignment = WD_TABLE_ALIGNMENT.CENTER
        else:
            self.addWarnParagraph(doc,"此处对应的sql未查询到内容，请注意检查")

    def scword(self):
        '''生成文档'''
        self.moveoldfile()  # 处理附件目录
        # sql查询数据
        getdata_dx = getsqldata()  # 获取sql的类对象
        sqltext = sqlstr(self.ny)  # 获取sql文本类对象
        all_text = get_ybwb.hqnr(get_ybwb())  # 获取到结构内容
        # 公共部分
        doc = self.doc
        # 一些样式
        style = doc.styles['Normal']
        style.font.name = u'Times New Roman'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        style.font.color.rgb = RGBColor(0, 0, 0)
        nf = int(self.ny[:4])
        yf = int(self.ny[-2:])
        end = calendar.monthrange(nf, yf)[1]
        table_style = "Table Grid"
        # 各个级别的字体对照大小
        ztdx_jb1 = 22
        ztdx_jb2 = 16
        ztdx_jb3 = 15
        ztdx_jb4 = 14
        ztdx_ym = 9
        ztdx_nr = 12
        ztdx_5h = 10.5

        # 获取当前时间
        curYear = time.localtime().tm_year
        curMonth = time.localtime().tm_mon
        curDay = time.localtime().tm_mday

        ################################## 添加文档页眉 ##################################
        header = doc.sections[0].header
        ymdl = header.paragraphs[0]
        run = ymdl.add_run()
        run.add_picture(os.path.join(self.root_dir, 'img/img_ym01.png'), width=Cm(2.5), height=Cm(0.8))
        run.add_picture(os.path.join(self.root_dir, 'img/img_ym02.png'), width=Cm(1.1), height=Cm(0.8))
        ymnr = '\t' + ' ' * 82 + '{nf}年数据资源运行维护管理项目'.format(nf=nf)
        run.add_text(ymnr)
        font = run.font
        font.size = Pt(ztdx_ym)
        run.underline = True
        ymdl.style = doc.styles["Header"]
        ymgs = ymdl.paragraph_format
        ymgs.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        ################################## 封面 ##################################
        paragraph = doc.add_paragraph()
        dbt = paragraph.add_run('任务运行情况月报\n{nf}年{yf}月'.format(nf=nf, yf=yf))
        dbt.bold = True
        dbt.font.name = u'黑体'
        dbt._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        # 字体对象
        btzt = dbt.font
        btzt.size = Pt(ztdx_jb1)  # 设置标题字体大小
        btzt.bold = True
        btzt.color.rgb = RGBColor(0, 0, 0)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format.space_after = Pt(ztdx_jb4 * 9)
        paragraph_format.space_before = Pt(ztdx_jb1 * 4)
        paragraph_format.line_spacing = Pt(ztdx_jb1 * 1.5)

        # 添加一个文档说明表格 table_style
        wdsmbg = doc.add_table(rows=7, cols=4, style=table_style)
        wdsmbg.alignment = WD_TABLE_ALIGNMENT.CENTER
        wdsmbg.cell(0, 0).text = '文档敏感性定义'
        wdsmbg.cell(0, 1).text = '敏感'
        wdsmbg.cell(0, 1).merge(wdsmbg.cell(0, 2)).merge(wdsmbg.cell(0, 3))  # 合并后面两个格子的对象

        wdsmbg.cell(1, 0).text = '编写人'
        wdsmbg.cell(1, 1).text = r'个人/组'
        wdsmbg.cell(1, 2).text = '编写日期'
        wdsmbg.cell(1, 3).text = r'YYYY-MM-DD'

        wdsmbg.cell(2, 0).text = '审核人'
        wdsmbg.cell(2, 1).text = '个人'
        wdsmbg.cell(2, 2).text = '审核日期'
        wdsmbg.cell(2, 3).text = r'YYYY-MM-DD'

        wdsmbg.cell(3, 0).text = '公开范围'
        wdsmbg.cell(3, 1).text = '国家税务总局、云平台管理项目组'
        wdsmbg.cell(3, 1).merge(wdsmbg.cell(3, 2)).merge(wdsmbg.cell(3, 3))

        wdsmbg.cell(4, 0).text = '建设单位'
        wdsmbg.cell(4, 1).text = '国家税务总局'
        wdsmbg.cell(4, 1).merge(wdsmbg.cell(4, 2)).merge(wdsmbg.cell(4, 3))

        wdsmbg.cell(5, 0).text = '承建单位'
        wdsmbg.cell(5, 1).text = '中国软件与技术服务股份有限公司'
        wdsmbg.cell(5, 1).merge(wdsmbg.cell(5, 2)).merge(wdsmbg.cell(5, 3))

        wdsmbg.cell(6, 0).text = '监理单位'
        wdsmbg.cell(6, 1).text = '北京赛迪工业和信息化工程监理中心有限公司'
        wdsmbg.cell(6, 1).merge(wdsmbg.cell(6, 2)).merge(wdsmbg.cell(6, 3))

        # 将表格内容设置为居中
        for r in range(7):
            for c in range(4):
                cell = wdsmbg.cell(r, c)
                pf = cell.paragraphs[0].paragraph_format
                pf.alignment = WD_TABLE_ALIGNMENT.CENTER
                pf.line_spacing = Pt(ztdx_5h * 3)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        doc.add_page_break()

        ################################## 目录部分 ##################################

        paragraph = doc.add_paragraph()
        xdbt = paragraph.add_run("目  录")
        xdbtzt = xdbt.font
        xdbtzt.name = u'黑体'
        xdbt._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
        xdbtzt.size = Pt(ztdx_jb1)
        xdbtzt.bold = True
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # # 段落前后的空行数量
        paragraph_format.space_after = Pt(ztdx_jb1 / 2)
        paragraph_format.space_before = Pt(ztdx_jb1 / 2)

        doc.add_page_break()

        ################################## 添加正文，这里的456三章其实结构完全相同，可以设计一个公共方法来写 ##################################

        for jb1 in all_text.keys():
            dldl = doc.add_heading(level=1)  # 添加标题
            pf = dldl.paragraph_format
            pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pf.space_before = Pt(17)
            pf.space_after = Pt(17)
            pf.line_spacing = Pt(ztdx_jb1 * 2.41)
            run = dldl.add_run(jb1)
            run.font.name = u'宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            dlbtzt = run.font
            dlbtzt.bold = True
            dlbtzt.size = Pt(ztdx_jb1)
            # 设置颜色，这两种都可以
            dlbtzt.color.rgb = RGBColor(0, 0, 0)

            # 处理第二级
            for jb2 in all_text.get(jb1):
                for jb2key in jb2.keys():
                    self.addTitle(doc, ztdx_jb2, jb2key, 1.73, 13, 2, u'黑体')
                    # 处理标准结构中特殊内容的  很神奇，这个jb2.get(jb2key) 和jb2.values() 居然拿到的内容不一样,因为python还会自动在外面套一层
                    if jb2key[:jb2key.index("\t")] == "1.1" or jb2key[:jb2key.index("\t")] == "1.2" or jb2key[
                                                                                                       :jb2key.index(
                                                                                                           "\t")] == "1.3":
                        for nr in jb2.get(jb2key):
                            self.addSpecialParagraph2(nr, ztdx_nr, doc, "、", "：")
                    elif jb2key[:jb2key.index("\t")] == "2.1":
                        for nr in jb2.get(jb2key):
                            self.addParagraph(nr, ztdx_nr, doc)
                            # 动态内容
                            data=getdata_dx.getzlzData(sqltext.get_tj2101_xzxm())
                            # print("11111",data)
                            nr=["{cjmc}新增{gs}生产项目空间，分别是{text};".format(cjmc=x[0],gs=x[1],text=x[2]) for x in data]
                            temp_nr = "{nf}年{yf}月，{nr}".format(nf=nf, yf=yf,nr=reduce(lambda a,b:a+b,nr))
                            self.addParagraph(temp_nr,ztdx_nr,doc)
                            temp_nr = "{nf}年{yf}月，各数据层的生产项目空间数量统计如下：".format(nf=nf, yf=yf)
                            self.addParagraph(temp_nr, ztdx_nr, doc)
                            data = getdata_dx.getzlzData(sqltext.get_b2101_cjxmsl())
                            temp_list = [len(data) + 1, '总计', self.jszs(data, 2), self.jszs(data, 3)]
                            data=list(data)
                            data.append(temp_list)
                            tr_text = ['序号', '数据层级', '项目数量', '比上月项目数增量']
                            tr_width = [1, 3, 3, 3]
                            self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                            temp_nr = "截止{nf}年{yf}月{end}日，各数据层需重点关注的生产项目空间的任务数量统计如下：".format(nf=nf, yf=yf, end=end)
                            self.addParagraph(temp_nr, ztdx_nr, doc)
                            data = getdata_dx.getzlzData(sqltext.get_b2102_zdxmrwtj())
                            temp_list = [len(data) + 1, '总计', self.jszs(data, 2), self.jszs(data, 3),self.jszs(data,4)]
                            data = list(data)
                            data.append(temp_list)
                            tr_text = ['序号', '数据层级', '项目数量', '任务数量', '比上月任务数增量']
                            tr_width = [1, 1.5, 1.5, 1.5, 4.5]
                            self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                            temp_nr = "（注：1、镜像层的项目空间下有大量历史任务，经过初步筛选，目前统计出共镜像层项目空间下有252613个任务是历史任务，本次统计全文的历史任务均不在统计范围内。"
                            self.addParagraph(temp_nr, ztdx_nr, doc)
                            temp_nr = "2、本月任务数量=上月任务数量+本月新创建任务数－本月下线任务数"
                            self.addParagraph(temp_nr, ztdx_nr, doc)
                            temp_nr = "3、重点关注：即梳理分析任务运行情况与调度情况时由中软公司主要负责的项目，具体包含镜像层的36个省局项目+1个总局项目，基础层与中间层非私有空间的项目，模型层由中软创建的50个项目。)"
                            self.addParagraph(temp_nr, ztdx_nr, doc)
                            data = getdata_dx.getzlzData(sqltext.get_b2103_zdxmxzrwtj())
                            temp_list = [len(data) + 1, '总计', self.jszs(data, 2), self.jszs(data, 3)]
                            data = list(data)
                            data.append(temp_list)
                            temp_nr = "{nf}年{yf}月云平台各生产项目空间新创建项目共计{xzrws}个任务。各数据层创建的任务数量统计如下：".format(nf=nf, yf=yf,
                                                                                                      xzrws=self.jszs(data, 2))
                            self.addParagraph(temp_nr, ztdx_nr, doc)
                            tr_text = ['序号', '数据层级', '新创建任务数量', '比上月新增任务数增量']
                            tr_width = [1, 2, 3, 4]
                            self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    elif jb2key[:jb2key.index("\t")] == "2.2":
                        for nr in jb2.get(jb2key):
                            self.addParagraph(nr, ztdx_nr, doc)
                        tr_text = ['序号', '数据层级', '总数', '正常调度', '手动调度', '暂停调度', '空跑调度']
                        tr_width = [1, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5]
                        data = getdata_dx.getzlzData(sqltext.get_b2201_fdtlxtj())
                        temp_list = [len(data) + 1, '总计', self.jszs(data, 2), self.jszs(data, 3),self.jszs(data, 4), self.jszs(data, 5),self.jszs(data, 6)]
                        data = list(data)
                        data.append(temp_list)
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    elif jb2key[:jb2key.index("\t")] == "2.3":
                        for nr in jb2.get(jb2key):
                            self.addParagraph(nr, ztdx_nr, doc)
                        tr_text = ['序号', '数据层级', '实例总数', '未运行', '运行中', '运行失败', '运行成功', '比上月实例数增量']
                        tr_width = [0.5, 1.5, 2, 1, 1.3, 1.2, 1.5, 2]
                        data = getdata_dx.getzlzData(sqltext.get_b2301_zcrwfyxzttj())
                        temp_list = [len(data) + 1, '总计', self.jszs(data, 2), self.jszs(data, 3), self.jszs(data, 4),
                                     self.jszs(data, 5), self.jszs(data, 6), self.jszs(data, 7)]
                        data = list(data)
                        data.append(temp_list)
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    elif jb2key[:jb2key.index("\t")] == "3.1":
                        data=getdata_dx.getzlzData(sqltext.get_tj3101_jxcxm())
                        temp_nr = "截至{nf}年{yf}月{end}日，镜像层的生产项目共有{zgs}个，包括36个国税项目（由于国地税已合并，目前不使用地税项目，仅使用国税项目）、1个快速加工项目、34个地税项目、36个准实时项目，以及实名办税、汇算清缴、国税总局等5个其他项目。在镜像层的项目中38个项目是正常运维使用，34个地税项目已停用，36个准实时项目由集成组每天巡检，全部任务正常运行完成，其余4个项目无专人维护。本次报告对{zdgs}个项目空间任务进行分析。".format(
                            nf=nf,yf=yf,end=end,zgs=data[0][0],zdgs=data[0][1]
                        )
                        self.addParagraph(temp_nr,ztdx_nr,doc)
                        for nr in jb2.get(jb2key):
                            # 最小级别的内容
                            self.addParagraph(nr, ztdx_nr, doc)
                        data = getdata_dx.getzlzData(sqltext.get_b3101_scrwsltj())
                        temp_list = [len(data) + 1, '总计', self.jszs(data, 2), self.jszs(data, 3)]
                        data = list(data)
                        data.append(temp_list)
                        tr_text = ['序号', '项目名称', '任务数量', '比上月任务增量']
                        tr_width = [1, 3, 2, 3]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        data = getdata_dx.getzlzData(sqltext.get_b3102_scxzrw())
                        temp_nr = "{nf}年{yf}月，镜像层生产项目空间新增任务{xzrws}个，新增任务按项目统计如下表所示：".format(nf=nf, yf=yf,
                                                                                            xzrws=data[-1][-1])
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        data = getdata_dx.getzlzData(sqltext.get_b3102_scxzrw())
                        temp_list = [len(data) + 1, '总计', self.jszs(data, 2)]
                        data = list(data)
                        data.append(temp_list)
                        tr_text = ['序号', '项目名称', '新创建任务数量']
                        tr_width = [1, 3, 3]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        self.addParagraph("新增任务说明：", ztdx_nr, doc)
                    elif jb2key[:jb2key.index("\t")] == "3.2":
                        for jb3 in jb2.get(jb2key):
                            for jb3key in jb3.keys():
                                self.addTitle(doc, ztdx_jb3, jb3key)
                                if jb3key[:jb3key.index("\t")] == "3.2.1":
                                    # 最小级别的内容
                                    data = getdata_dx.getzlzData(sqltext.get_b3201_ddlxtjrws())
                                    temp_nr = "截至{nf}年{yf}月{end}日，镜像层生产项目空间的任务总数为{zs}个，按任务调度类型统计如下表所示：".format(nf=nf,
                                                                                                               yf=yf,
                                                                                                               end=end,
                                                                                                               zs=self.jszs(
                                                                                                                   data,
                                                                                                                   2))
                                    self.addParagraph(temp_nr, ztdx_nr, doc)
                                    tr_text = ['序号', '任务调度类型', '任务数量']
                                    tr_width = [1, 3, 3]
                                    self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                                if jb3key[:jb3key.index("\t")] == "3.2.2":
                                    # 最小级别的内容
                                    data = getdata_dx.getzlzData(sqltext.get_b3202_zcrwalxtj())
                                    temp_nr = "截至{nf}年{yf}月{end}日，正常调度的任务数量是{zs}个，按任务类型统计如下表所示：".format(nf=nf, yf=yf,
                                                                                                        end=end,
                                                                                                        zs=self.jszs(
                                                                                                            data, 2))
                                    self.addParagraph(temp_nr, ztdx_nr, doc)
                                    tr_text = ['序号', '任务类型', '任务数量']
                                    tr_width = [1, 3, 3]
                                    self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    elif jb2key[:jb2key.index("\t")] == "3.3":
                        for jb3 in jb2.get(jb2key):
                            for jb3key in jb3.keys():
                                self.addTitle(doc, ztdx_jb3, jb3key)
                                if jb3key[:jb3key.index("\t")] == "3.3.1":
                                    # 最小级别的内容
                                    data = getdata_dx.getzlzData(sqltext.get_b3301_rwyxzttj())
                                    temp_nr = "据{curyear}年{curmonth}月{curday}日查询结果，{nf}年{yf}月1日至{nf}年{yf}月{end}日，镜像层生产项目空间中正常调度的任务，除虚节点任务（由于虚节点的特殊性质见第1章1.2节虚节点概念介绍，无节点代码，因此不在任务运行状态统计范围内，后面章节同理。）总共运行产生{slzs}个实例。任务运行状态如下表所示：" \
                                        .format(curyear=curYear, curmonth=curMonth, curday=curDay, nf=nf, yf=yf,
                                                end=end, slzs=self.jszs(data, 2))
                                    self.addParagraph(temp_nr, ztdx_nr, doc)
                                    tr_text = ['序号', '任务运行状态', '实例数']
                                    tr_width = [1, 4, 4]
                                    self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                                    data = getdata_dx.getzlzData(sqltext.get_b3302_sbrwqk())
                                    if data:
                                        temp_nr = "{yf}月镜像层生产项目中，仅SC_JX_KSJG项目空间下{sbgs}个任务失败，如下表所示：".format(yf=yf,
                                                                                                            sbgs=
                                                                                                            data[-1][0])
                                        self.addParagraph(temp_nr, ztdx_nr, doc)
                                        tr_text = ['序号', '项目名称', '任务名称', '失败次数']
                                        tr_width = [1, 3, 4, 2]
                                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                                    else:
                                        self.addWarnParagraph(doc,"get_b3302_sbrwqk未获取到内容，请注意检查！！！")
                                    data = getdata_dx.getzlzData(sqltext.get_b3303_wyxrwqk())
                                    if data:
                                        temp_nr = "{yf}月镜像层生产项目中，未运行的任务有{wyxgs}个，如下表所示：".format(yf=yf,
                                                                                                wyxgs=data[-1][0])
                                        self.addParagraph(temp_nr, ztdx_nr, doc)
                                        tr_text = ['序号', '项目名称', '任务名称', '未运行次数']
                                        tr_width = [1, 3, 4, 2]
                                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                                        temp_nr = "任务未运行原因："
                                        self.addParagraph(temp_nr, ztdx_nr, doc)
                                    else:
                                        self.addWarnParagraph(doc, "get_b3303_wyxrwqk未获取到内容，请注意检查！！！")
                    elif jb2key[:jb2key.index("\t")] == "4.1":
                        # 最小级别的内容
                        data = getdata_dx.getzlzData(sqltext.get_b4101_jccscxmkj_info())
                        temp_nr = "截至{nf}年{yf}月{end}日，基础层共有{gs}个生产项目空间，其中有8个私有空间。私有空间为各省提供自主开发使用，不做日常运维，因此不做展开分析。基础层的生产项目空间如下表所示。" \
                            .format(nf=nf, yf=yf, end=end, gs=data[-1][0])
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        temp_nr = '（注：基础层生产项目统计口径：项目空间名称以“SC_JC_"开头，且不包含"SC_JC_ZRRYH"）'
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        tr_text = ['序号', '项目名称', '项目中文名称']
                        tr_width = [1, 4, 4]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        data = getdata_dx.getzlzData(sqltext.get_b4102_jccrwtj())
                        temp_nr = "现对基础层的SC_JC_GSZJ、SC_JC_GSAQ，共2个项目空间进行任务数量统计，共计{rwgs}个任务。这些项目的创建厂商均为中软。".format(
                            rwgs=self.jszs(data, 2))
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        tr_text = ['序号', '项目名称', '任务数量', '比上月任务增量']
                        tr_width = [1, 3, 3, 3]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        data = getdata_dx.getzlzData(sqltext.get_b4103_jjcxzrwtj())
                        nr = reduce(
                            lambda a, b: "{xmmc}项目空间新增{rwgs}个任务，{xmmc1}项目空间新增{rwgs1}个任务。".format(xmmc=a[1], rwgs=a[2],
                                                                                                 xmmc1=b[1],
                                                                                                 rwgs1=b[2]), data)
                        temp_nr = "{nf}年{yf}月，基础层生产项目空间共新增任务{xzzs}个，{pjnr}".format(nf=nf, yf=yf,
                                                                                   xzzs=self.jszs(data, 2), pjnr=nr)
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        temp_list = [len(data) + 1, '总计', self.jszs(data, 2)]
                        data = list(data)
                        data.append(temp_list)
                        tr_text = ['序号', '项目名称', '新创建任务数量']
                        tr_width = [1, 3, 3]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        self.addParagraph("新增任务说明：", ztdx_nr, doc)
                    elif jb2key[:jb2key.index("\t")] == "4.2":
                        for jb3 in jb2.get(jb2key):
                            for jb3key in jb3.keys():
                                self.addTitle(doc, ztdx_jb3, jb3key)
                                if jb3key[:jb3key.index("\t")] == "4.2.1":
                                    # 最小级别的内容
                                    data = getdata_dx.getzlzData(sqltext.get_b4201_ddlxtj())
                                    temp_nr = "截至{nf}年{yf}月{end}日，基础层生产项目空间的任务总数为{rwzs}个，按任务调度类型统计如下表所示。".format(nf=nf,
                                                                                                                 yf=yf,
                                                                                                                 end=end,
                                                                                                                 rwzs=self.jszs(
                                                                                                                     data,
                                                                                                                     2))
                                    self.addParagraph(temp_nr, ztdx_nr, doc)
                                    tr_text = ['序号', '任务调度类型', '任务数量']
                                    tr_width = [1, 4, 4]
                                    self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                                if jb3key[:jb3key.index("\t")] == "4.2.2":
                                    # 最小级别的内容
                                    data = getdata_dx.getzlzData(sqltext.get_b4202_zcrwalxtj())
                                    temp_nr = "截至{nf}年{yf}月{end}日，正常调度的任务数量是{rwzs}个，按任务类型统计如下表所示。".format(nf=nf, yf=yf,
                                                                                                          end=end,
                                                                                                          rwzs=self.jszs(
                                                                                                              data, 2))
                                    self.addParagraph(temp_nr, ztdx_nr, doc)
                                    tr_text = ['序号', '任务调度类型', '任务数量']
                                    tr_width = [1, 4, 4]
                                    self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    elif jb2key[:jb2key.index("\t")] == "4.3":
                        data = getdata_dx.getzlzData(sqltext.get_b4301_rwyxzttj())
                        temp_nr = "根据{curyear}年{curmonth}月{curday}日查询数据， {nf}年{yf}月1日至{nf}年{yf}月{end}日，基础层生产项目空间中正常调度的任务，除虚节点任务，运行产生{slgs}个实例。任务运行状态如下表所示。" \
                            .format(curyear=curYear, curmonth=curMonth, curday=curDay, nf=nf, yf=yf, end=end,
                                    slgs=self.jszs(data, 2))
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        tr_text = ['序号', '任务运行状态', '实例数']
                        tr_width = [1, 4, 4]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        data = getdata_dx.getzlzData(sqltext.get_b4302_wyxrwmx())
                        temp_nr = "{nf}年{yf}月基础层生产项目中，未运行的任务有{wyxgs}个，如下表所示：".format(nf=nf, yf=yf, wyxgs=len(data))
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        tr_text = ['序号', '项目名称', '任务名称', '未运行次数']
                        tr_width = [1, 2, 4, 2]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        self.addParagraph("经分析，", ztdx_nr, doc)
                        data = getdata_dx.getzlzData(sqltext.get_b4303_sbrwmx())
                        temp_nr = "{nf}年{yf}月基础层生产项目中，失败的任务有{wyxgs}个，如下表所示：".format(nf=nf, yf=yf, wyxgs=len(data))
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        tr_text = ['序号', '项目名称', '任务名称', '失败次数']
                        tr_width = [1, 2, 4, 2]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        self.addParagraph("经分析，", ztdx_nr, doc)
                    elif jb2key[:jb2key.index("\t")] == "4.4":
                        for nr in jb2.get(jb2key):
                            self.addParagraph(nr, ztdx_nr, doc)
                        data = getdata_dx.getzlzData(sqltext.get_b4401_xhzyzg())
                        tr_text = ['排名', '项目名称', r'任务名称/任务信息', r'计算资源（核*运行时长）', '实例开始时间', '实例结束时间', '实例运行时长']
                        tr_width = [0.5, 1.5, 2, 1.5, 1.5, 1.5, 1.5]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    elif jb2key[:jb2key.index("\t")] == "4.5":
                        data = getdata_dx.getzlzData(sqltext.get_b4501_rwyxsc())
                        temp_nr = "{nf}年{yf}月1日至{nf}年{yf}月{end}日，基础层生产项目空间的任务运行时长最长的前10个任务如下表所示。".format(nf=nf, yf=yf,
                                                                                                         end=end)
                        self.addParagraph(nr, ztdx_nr, doc)
                        tr_text = ['排名', '项目名称', r'任务名称/任务信息', '任务开始时间', '任务结束时间', '任务运行时长', '日分区']
                        tr_width = [0.5, 1.5, 2, 1.5, 1.5, 1.5, 1.5]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    # -----------------------------中间层------------------------------------ #
                    elif jb2key[:jb2key.index("\t")] == "5.1":
                        # 最小级别的内容
                        data = getdata_dx.getzlzData(sqltext.get_b5101_jccscxmkj_info())
                        temp_nr = "截至{nf}年{yf}月{end}日，中间层共有{gs}个生产项目空间，其中有8个私有空间。私有空间为各省提供自主开发使用，不做日常运维，因此不做展开分析。中间层的生产项目空间如下表所示。" \
                            .format(nf=nf, yf=yf, end=end, gs=data[-1][0])
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        temp_nr = '（注：中间层生产项目统计口径：项目空间名称以“SC _ZJ_"开头，且不包含"SC_ZJ_ZRRYH"）'
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        tr_text = ['序号', '项目名称', '项目中文名称']
                        tr_width = [1, 4, 4]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        data = getdata_dx.getzlzData(sqltext.get_b5102_jccrwtj())
                        temp_nr = "现对中间层的SC_ZJ_GSZJ、SC_ZJ_GSZJ_2共2个生产项目空间进行任务数量统计，共计{rwgs}个任务。这些项目的创建厂商均为中软。按项目统计的任务数量如下表所示：".format(
                            rwgs=self.jszs(data, 2))
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        tr_text = ['序号', '项目名称', '任务数量', '比上月任务增量']
                        tr_width = [1, 3, 3, 3]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        data = getdata_dx.getzlzData(sqltext.get_b5103_jjcxzrwtj())
                        nr = reduce(
                            lambda a, b: "{xmmc}项目空间新增{rwgs}个任务，{xmmc1}项目空间新增{rwgs1}个任务。".format(xmmc=a[1], rwgs=a[2],
                                                                                                 xmmc1=b[1],
                                                                                                 rwgs1=b[2]), data)
                        temp_nr = "{nf}年{yf}月，中间层生产项目空间共新增任务{xzzs}个，{pjnr}".format(nf=nf, yf=yf,
                                                                                   xzzs=self.jszs(data, 2), pjnr=nr)
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        temp_list = [len(data) + 1, '总计', self.jszs(data, 2)]
                        data = list(data)
                        data.append(temp_list)
                        tr_text = ['序号', '项目名称', '新创建任务数量']
                        tr_width = [1, 3, 3]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        self.addParagraph("新增任务说明：", ztdx_nr, doc)
                    elif jb2key[:jb2key.index("\t")] == "5.2":
                        for jb3 in jb2.get(jb2key):
                            for jb3key in jb3.keys():
                                self.addTitle(doc, ztdx_jb3, jb3key)
                                if jb3key[:jb3key.index("\t")] == "5.2.1":
                                    # 最小级别的内容
                                    data = getdata_dx.getzlzData(sqltext.get_b5201_ddlxtj())
                                    temp_nr = "截至{nf}年{yf}月{end}日，中间层生产项目空间的任务总数为{rwzs}个，按任务调度类型统计如下表所示。".format(nf=nf,
                                                                                                                 yf=yf,
                                                                                                                 end=end,
                                                                                                                 rwzs=self.jszs(
                                                                                                                     data,
                                                                                                                     2))
                                    self.addParagraph(temp_nr, ztdx_nr, doc)
                                    tr_text = ['序号', '任务调度类型', '任务数量']
                                    tr_width = [1, 4, 4]
                                    self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                                if jb3key[:jb3key.index("\t")] == "5.2.2":
                                    # 最小级别的内容
                                    data = getdata_dx.getzlzData(sqltext.get_b5202_zcrwalxtj())
                                    temp_nr = "截至{nf}年{yf}月{end}日，正常调度的任务数量是{rwzs}个，按任务类型统计如下表所示。".format(nf=nf, yf=yf,
                                                                                                          end=end,
                                                                                                          rwzs=self.jszs(
                                                                                                              data, 2))
                                    self.addParagraph(temp_nr, ztdx_nr, doc)
                                    tr_text = ['序号', '任务调度类型', '任务数量']
                                    tr_width = [1, 4, 4]
                                    self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    elif jb2key[:jb2key.index("\t")] == "5.3":
                        data = getdata_dx.getzlzData(sqltext.get_b5301_rwyxzttj())
                        temp_nr = "根据{curyear}年{curmonth}月{curday}日查询数据， {nf}年{yf}月1日至{nf}年{yf}月{end}日，中间层生产项目空间中正常调度的任务，除虚节点任务，运行产生{slgs}个实例。任务运行状态如下表所示。" \
                            .format(curyear=curYear, curmonth=curMonth, curday=curDay, nf=nf, yf=yf, end=end,
                                    slgs=self.jszs(data, 2))
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        tr_text = ['序号', '任务运行状态', '实例数']
                        tr_width = [1, 4, 4]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        data = getdata_dx.getzlzData(sqltext.get_b5303_sbrwmx())
                        temp_nr = "{nf}年{yf}月中间层生产项目中，失败的任务有{wyxgs}个，如下表所示：".format(nf=nf, yf=yf, wyxgs=len(data))
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        tr_text = ['序号', '项目名称', '任务名称', '失败次数']
                        tr_width = [1, 2, 4, 2]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        self.addParagraph("经分析，", ztdx_nr, doc)
                    elif jb2key[:jb2key.index("\t")] == "5.4":
                        for nr in jb2.get(jb2key):
                            self.addParagraph(nr, ztdx_nr, doc)
                        data = getdata_dx.getzlzData(sqltext.get_b5401_xhzyzg())
                        tr_text = ['排名', '项目名称', '任务名称/任务信息', '计算资源（核*运行时长）', '实例开始时间', '实例结束时间', '实例运行时长']
                        tr_width = [0.5, 1.5, 2, 1.5, 1.5, 1.5, 1.5]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    elif jb2key[:jb2key.index("\t")] == "5.5":
                        data = getdata_dx.getzlzData(sqltext.get_b4501_rwyxsc())
                        temp_nr = "{nf}年{yf}月1日至{nf}年{yf}月{end}日，中间层生产项目空间的任务运行时长最长的前10个任务如下表所示。".format(nf=nf, yf=yf,
                                                                                                         end=end)
                        self.addParagraph(nr, ztdx_nr, doc)
                        tr_text = ['排名', '项目名称', '任务名称/任务信息', '任务开始时间', '任务结束时间', '任务运行时长', '日分区']
                        tr_width = [0.5, 1.5, 2, 1.5, 1.5, 1.5, 1.5]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    # -----------------------------模型层------------------------------------ #
                    elif jb2key[:jb2key.index("\t")] == "6.1":
                        # 最小级别的内容
                        data = getdata_dx.getzlzData(sqltext.get_b6101_jccscxmkj_info())
                        temp_nr = "截至{nf}年{yf}月{end}日，模型层目前梳理出属于中软厂商创建的，且仍然在使用的生产项目空间有{xmgs}个，总共{rwgs}个任务。各生产项目空间包含的任务数量如下表所示。" \
                            .format(nf=nf, yf=yf, end=end, xmgs=len(data), rwgs=self.jszs(data, 2))
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        temp_nr = '（注：模型层项目数量统计口径：由模型层小组提供具体项目清册）'
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        temp_list = [len(data) + 1, '总计', self.jszs(data, 2), self.jszs(data, 3)]
                        data = list(data)
                        data.append(temp_list)
                        tr_text = ['序号', '项目名称', '任务数量', '比上月任务增量']
                        tr_width = [1, 3, 2, 4]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        data = getdata_dx.getzlzData(sqltext.get_b6102_jccrwtj())
                        temp_nr = "模型层生产项目在{nf}年{yf}月新增了{rwgs}个任务，各项目新增任务数量如下表所示。".format(nf=nf, yf=yf,
                                                                                          rwgs=self.jszs(data, 2))
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        temp_list = [len(data) + 1, '总计', self.jszs(data, 2)]
                        data = list(data)
                        data.append(temp_list)
                        tr_text = ['序号', '项目名称', '新创建任务数量']
                        tr_width = [1, 4, 3]

                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        data = getdata_dx.getzlzData(sqltext.get_b6103_jjcxzrwtj())
                        nr = reduce(
                            lambda a, b: "{xmmc}项目空间新增{rwgs}个任务，{xmmc1}项目空间新增{rwgs1}个任务。".format(xmmc=a[1], rwgs=a[2],
                                                                                                 xmmc1=b[1],
                                                                                                 rwgs1=b[2]), data)
                        temp_nr = "{nf}年{yf}月，模型层生产项目空间共新增任务{xzzs}个，{pjnr}".format(nf=nf, yf=yf,
                                                                                   xzzs=self.jszs(data, 2), pjnr=nr)
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        self.addParagraph("新增任务说明：", ztdx_nr, doc)
                    elif jb2key[:jb2key.index("\t")] == "6.2":
                        for jb3 in jb2.get(jb2key):
                            for jb3key in jb3.keys():
                                self.addTitle(doc, ztdx_jb3, jb3key)
                                if jb3key[:jb3key.index("\t")] == "6.2.1":
                                    # 最小级别的内容
                                    data = getdata_dx.getzlzData(sqltext.get_b6201_ddlxtj())
                                    temp_nr = "截至{nf}年{yf}月{end}日，模型层生产项目空间的任务总数为{rwzs}个，按任务调度类型统计如下表所示。".format(nf=nf,
                                                                                                                 yf=yf,
                                                                                                                 end=end,
                                                                                                                 rwzs=self.jszs(
                                                                                                                     data,
                                                                                                                     2))
                                    self.addParagraph(temp_nr, ztdx_nr, doc)
                                    tr_text = ['序号', '任务调度类型', '任务数量']
                                    tr_width = [1, 4, 4]
                                    self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                                if jb3key[:jb3key.index("\t")] == "6.2.2":
                                    # 最小级别的内容
                                    data = getdata_dx.getzlzData(sqltext.get_b6202_zcrwalxtj())
                                    temp_nr = "截至{nf}年{yf}月{end}日，正常调度的任务数量是{rwzs}个，按任务类型统计如下表所示。".format(nf=nf, yf=yf,
                                                                                                          end=end,
                                                                                                          rwzs=self.jszs(
                                                                                                              data, 2))
                                    self.addParagraph(temp_nr, ztdx_nr, doc)
                                    tr_text = ['序号', '任务调度类型', '任务数量']
                                    tr_width = [1, 4, 4]
                                    self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    elif jb2key[:jb2key.index("\t")] == "6.3":
                        data = getdata_dx.getzlzData(sqltext.get_b6301_rwyxzttj())
                        temp_nr = "根据{curyear}年{curmonth}月{curday}日查询数据， {nf}年{yf}月1日至{nf}年{yf}月{end}日，模型层生产项目空间中正常调度的任务，除虚节点任务，运行产生{slgs}个实例。任务运行状态如下表所示。" \
                            .format(curyear=curYear, curmonth=curMonth, curday=curDay, nf=nf, yf=yf, end=end,
                                    slgs=self.jszs(data, 2))
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        tr_text = ['序号', '任务运行状态', '实例数']
                        tr_width = [1, 4, 4]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        data = getdata_dx.getzlzData(sqltext.get_b6302_sbrwmx())
                        temp_nr = "{nf}年{yf}月模型层生产项目中，失败的任务有{wyxgs}个，如下表所示：".format(nf=nf, yf=yf, wyxgs=len(data))
                        self.addParagraph(temp_nr, ztdx_nr, doc)
                        tr_text = ['序号', '项目名称', '任务名称', '失败次数']
                        tr_width = [1, 2, 4, 2]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                        self.addParagraph("经分析，", ztdx_nr, doc)
                    elif jb2key[:jb2key.index("\t")] == "6.4":
                        for nr in jb2.get(jb2key):
                            self.addParagraph(nr, ztdx_nr, doc)
                        data = getdata_dx.getzlzData(sqltext.get_b6401_xhzyzg())
                        tr_text = ['排名', '项目名称', '任务名称/任务信息', '计算资源（核*运行时长）', '实例开始时间', '实例结束时间', '实例运行时长']
                        tr_width = [0.5, 1.5, 2, 1.5, 1.5, 1.5, 1.5]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    elif jb2key[:jb2key.index("\t")] == "6.5":
                        data = getdata_dx.getzlzData(sqltext.get_b4501_rwyxsc())
                        temp_nr = "{nf}年{yf}月1日至{nf}年{yf}月{end}日，模型层生产项目空间的任务运行时长最长的前10个任务如下表所示。".format(nf=nf, yf=yf,
                                                                                                         end=end)
                        self.addParagraph(nr, ztdx_nr, doc)
                        tr_text = ['排名', '项目名称', '任务名称/任务信息', '任务开始时间', '任务结束时间', '任务运行时长', '日分区']
                        tr_width = [0.5, 1.5, 2, 1.5, 1.5, 1.5, 1.5]
                        self.addTable(doc, data, table_style, tr_text, tr_width, ztdx_5h)
                    else:
                        for nr in jb2.get(jb2key):
                            self.addParagraph(nr, ztdx_nr, doc)
            if jb1[:jb1.index("\t")] != "第7章":
                doc.add_page_break()  # 除最后一章，每一章完成后换页
        doc.save(self.wdmc)