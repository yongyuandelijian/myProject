# 功能描述： 1 从word中批量读取所有表格， 2 按照对应关系将表的信息和对应列的信息匹配成一行，存入数据库
# aaa 20201106
from docx import Document

class GetWordTable(object):
    def getalltable(self,wordname):
        '''获取word中所有表格'''
        try:
            doc = Document(wordname)
        except Exception as e:
            msg="读取word失败，请检查文件是否存在>>>"+e
        else:
            msg="读取{wordname}成功".format(wordname=wordname)
        finally:
            print(msg)

        # =================从对象操作，先获取每个heading中有多少表格=======================#
        tables=doc.tables
        paragraphs=doc.paragraphs
        for dl in paragraphs:
            print(dl.text)

        # doc.add_heading('标题',0)
        # headings=doc.paragraphs
        # print(len(headings))
        print("共有表格：",len(tables))


    def getData(self,wordname):
        """获取表格上的字还有对应的表格"""
        doc=Document(wordname)
        print(len(doc.tables))

        print(len(doc.paragraphs))
        for dl in doc.paragraphs:
            print(dl.text)


        # print(len(doc.sections))
        # for pd in doc.sections:
        #     print(pd)
        # print(len(doc.styles))
        # for ys in doc.styles:
        #     print(ys)




if __name__ == '__main__':
    wordname="../test.docx"
    # GetWordTable.getalltable(GetWordTable,wordname)
    GetWordTable.getData(GetWordTable,wordname)